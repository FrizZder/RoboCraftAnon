import re
import os
from pathlib import Path
from natasha import Segmenter, MorphVocab, NewsEmbedding, NewsNERTagger, Doc
import docx
import pdfplumber
import json
import sys
import base64

segmenter = Segmenter()
morph_vocab = MorphVocab()
emb = NewsEmbedding()
ner_tagger = NewsNERTagger(emb)

labels = {
    "PER": "FULL_NAME",
    "LOC": "LOCATION",
    "ORG": "ORG",
    "DATE": "DATE",
    "TIME": "TIME",
    "MISC": "ENTITY"
}

def mask_entities(text):
    doc = Doc(text)
    doc.segment(segmenter)
    doc.tag_ner(ner_tagger)
    used = set()
    counts = {}
    for span in doc.spans:
        span.normalize(morph_vocab)
        lab = labels.get(span.type, span.type)
        counts[lab] = counts.get(lab, 0) + 1
        placeholder = f"<{lab}_{counts[lab]}>"
        if span.text not in used:
            text = text.replace(span.text, placeholder)
            used.add(span.text)
    return text

def mask_patterns(text):
    patterns = [
        (r'\b[\w\.-]+@[\w\.-]+\.\w+\b', '<EMAIL>'),
        (r'(?<!\d)(\+7|8)?[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}(?!\d)', '<PHONE>'),
        (r'\bИИН[\s:—-]*\d{12}\b', '<IIN>'),
        (r'\b\d{12}\b', '<IIN>'),
        (r'ID\s*№?\s*\d{6,}', '<PASSPORT>'),
        (r'паспорт.*?серии\s*\d+\s*№\s*\d+', '<PASSPORT>', re.IGNORECASE),
        (r'дом\s*\d+[\w/]*[,]?\s*кв\.?\s*\d+', '<ADDRESS>', re.IGNORECASE),
        (r'(улиц[аы]|проспект|переулок|проезд|шоссе)\s+[А-Яа-яёЁ\s\-]+', '<STREET>', re.IGNORECASE),
        (r'\d{6}(?=\D)', '<POSTAL_CODE>'),
        (r'@\w+', '<SOCIAL>'),
        (r'(https?://\S+|www\.\S+)', '<URL>'),
        (r'hh\.ru/resume/\S+', '<RESUME>'),
        (r'docs\.google\.com\S*', '<GOOGLE_DOC>'),
        (r'\b(?:\d[ -]*?){13,19}\b', '<CARD>'),
        (r'\b[АВЕКМНОРСТУХ]{1}\d{3}[АВЕКМНОРСТУХ]{2}\s?\d{2,4}\b', '<CAR>'),
        (r'ВУ\s*№?\s*\d{6,}', '<DRIVER_LICENSE>')
    ]
    for pat in patterns:
        if len(pat) == 2:
            text = re.sub(pat[0], pat[1], text)
        else:
            text = re.sub(pat[0], pat[1], text, flags=pat[2])
    return text

def anonymize(text):
    text = mask_patterns(text)
    text = mask_entities(text)
    return text

def read_txt(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def write_txt(path, text):
    with open(path, 'w', encoding='utf-8') as f:
        f.write(text)

def read_docx(path):
    doc = docx.Document(path)
    return "\n".join([para.text for para in doc.paragraphs])

def write_docx(path, text):
    doc = docx.Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.save(path)

def read_pdf(path):
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + '\n'
    return text

def process_file(filepath):
    try:
        ext = filepath.suffix.lower()
        output_path = filepath.parent / (filepath.stem + '_anonymized' + filepath.suffix)

        if ext == '.txt':
            text = read_txt(filepath)
            result = anonymize(text)
            write_txt(output_path, result)
            return {
                'success': True,
                'original_text': text,
                'anonymized_text': result,
                'output_path': str(output_path)
            }
        elif ext == '.docx':
            text = read_docx(filepath)
            result = anonymize(text)
            write_docx(output_path, result)
            return {
                'success': True,
                'original_text': text,
                'anonymized_text': result,
                'output_path': str(output_path)
            }
        elif ext == '.pdf':
            text = read_pdf(filepath)
            result = anonymize(text)
            txt_output = output_path.with_suffix('.txt')
            write_txt(txt_output, result)
            return {
                'success': True,
                'original_text': text,
                'anonymized_text': result,
                'output_path': str(txt_output)
            }
        else:
            return {
                'success': False,
                'error': f"Неподдерживаемый формат файла: {ext}"
            }
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

if __name__ == "__main__":
    try:
        input_data = json.loads(sys.stdin.readline())
        file_path = input_data.get('filePath')
        
        if not file_path:
            print(json.dumps({
                'success': False,
                'error': 'No file path provided'
            }))
            sys.exit(1)
        
        path = Path(file_path)
        ext = path.suffix.lower()
        if ext not in ['.txt', '.docx', '.pdf']:
            print(json.dumps({
                'success': False,
                'error': 'Unsupported file format'
            }))
            sys.exit(1)
        
        result = process_file(path)
        if result.get('success'):
            print(json.dumps({
                'success': True,
                'data': {
                    'original_text': result['original_text'],
                    'anonymized_text': result['anonymized_text'],
                    'output_path': result['output_path'],
                    'ext': ext
                }
            }))
        else:
            print(json.dumps({
                'success': False,
                'error': result.get('error', 'Unknown error')
            }))
        
    except Exception as e:
        print(json.dumps({
            'success': False,
            'error': str(e)
        }))
        sys.exit(1)
